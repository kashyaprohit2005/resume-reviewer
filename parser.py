import re
import io
from typing import Dict, Any, List

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None


class ResumeParser:
    """Document Parser to extract raw text and clean structure from PDF, DOCX, and TXT files."""

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text from a PDF file stream."""
        text = ""
        if pypdf:
            try:
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            except Exception as e:
                text = f"[PDF Parsing Error: {str(e)}]"
        else:
            # Fallback simple string decoding if pypdf missing
            text = file_bytes.decode('utf-8', errors='ignore')
        return text

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from a DOCX file stream."""
        text = ""
        if docx:
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs:
                    text += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        text += row_text + "\n"
            except Exception as e:
                text = f"[DOCX Parsing Error: {str(e)}]"
        else:
            text = file_bytes.decode('utf-8', errors='ignore')
        return text

    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> str:
        """Extract text from TXT file stream."""
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            return file_bytes.decode('latin-1', errors='ignore')

    @classmethod
    def parse_file(cls, file_bytes: bytes, filename: str) -> str:
        """Main dispatcher based on file extension."""
        ext = filename.lower().split('.')[-1]
        if ext == 'pdf':
            return cls.extract_text_from_pdf(file_bytes)
        elif ext in ['docx', 'doc']:
            return cls.extract_text_from_docx(file_bytes)
        else:
            return cls.extract_text_from_txt(file_bytes)

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize extracted text."""
        # Replace non-ascii quotes and bullets
        text = re.sub(r'[\u201c\u201d]', '"', text)
        text = re.sub(r'[\u2018\u2019]', "'", text)
        text = re.sub(r'[\u2022\u2023\u25e6\u2043\u2219]', '-', text)
        # Remove excessive blank lines and trailing spaces
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)

    @classmethod
    def extract_metadata(cls, text: str) -> Dict[str, Any]:
        """Extract contact info, email, phone, links, and estimate section headers."""
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        phone_pattern = r'\(?\+?\d{1,3}\)?[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}'
        linkedin_pattern = r'linkedin\.com/in/[a-zA-Z0-9_-]+'
        github_pattern = r'github\.com/[a-zA-Z0-9_-]+'

        emails = re.findall(email_pattern, text)
        phones = re.findall(phone_pattern, text)
        linkedins = re.findall(linkedin_pattern, text, re.IGNORECASE)
        githubs = re.findall(github_pattern, text, re.IGNORECASE)

        # Detect sections present
        section_headers = ['education', 'experience', 'work experience', 'projects', 'skills', 'certifications', 'summary', 'objective', 'publications', 'achievements']
        found_sections = []
        text_lower = text.lower()
        for header in section_headers:
            if re.search(r'\b' + header + r'\b', text_lower):
                found_sections.append(header.title())

        # Word count and sentence count
        words = re.findall(r'\w+', text)
        word_count = len(words)

        return {
            "email": emails[0] if emails else None,
            "phone": phones[0] if phones else None,
            "linkedin": f"https://{linkedins[0]}" if linkedins else None,
            "github": f"https://{githubs[0]}" if githubs else None,
            "word_count": word_count,
            "detected_sections": found_sections
        }
